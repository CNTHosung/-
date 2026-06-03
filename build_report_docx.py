# -*- coding: utf-8 -*-
"""REPORT 내용을 Word(.docx) 보고서로 생성. (python-docx)"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

HERE = os.path.dirname(os.path.abspath(__file__))
KOR = "Malgun Gothic"


def set_kor(style):
    rpr = style.element.get_or_add_rPr()
    rf = rpr.get_or_add_rFonts()
    for a in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rf.set(qn(a), KOR)


doc = Document()

# --- 기본/제목 스타일 한글 폰트 ---
normal = doc.styles["Normal"]
normal.font.name = KOR
normal.font.size = Pt(10.5)
set_kor(normal)
for sname in ("Heading 1", "Heading 2", "Heading 3", "Title"):
    try:
        set_kor(doc.styles[sname])
    except KeyError:
        pass

# 페이지 여백
for s in doc.sections:
    s.top_margin = s.bottom_margin = Inches(1)
    s.left_margin = s.right_margin = Inches(1)


def h(text, level=1):
    p = doc.add_heading(text, level=level)
    return p


def para(text, bold=False, italic=False, size=None, align=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if size:
        r.font.size = Pt(size)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def bullet(text):
    doc.add_paragraph(text, style="List Bullet")


def code_block(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.2)
    return p


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, htext in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        run = c.paragraphs[0].add_run(htext)
        run.bold = True
        run.font.size = Pt(9.5)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(9.5)
    return t


def image(path, width=6.4, caption=None):
    if os.path.exists(os.path.join(HERE, path)):
        doc.add_picture(os.path.join(HERE, path), width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            para(caption, italic=True, size=9, align="center")
    else:
        para(f"[그림 없음: {path}]", italic=True, size=9)


# ===========================================================================
# 표지
# ===========================================================================
title = doc.add_heading("SNN 기반 Light-Intensity → Frequency\n변환 소자 학습 보고서", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
para("작성일: 2026-06-03   |   프로젝트: C:\\Claude\\snn_light2freq", align="center", size=10)
para("프레임워크: snnTorch 0.9.4 (PyTorch 2.5.1+cpu)", align="center", size=10)
doc.add_paragraph()

# 1
h("1. 개요 및 목적", 1)
para("본 프로젝트는 광응답 소자(photoresponsive device)를 스파이킹 뉴런 하드웨어로 "
     "간주하고, 소자의 물리적 \"광세기(light intensity) → 스파이크 주파수(frequency)\" 변환 "
     "특성을 실측 데이터로부터 인코더에 직접 반영하여, MNIST 손글씨 숫자를 스파이킹 신경망"
     "(Spiking Neural Network, SNN)으로 분류하는 것을 목표로 한다.")
para("실제 광응답 소자는 입사 광세기 Φ가 커질수록 광전류 I_ph가 증가하고, 이 전류가 막전위를 "
     "충전시켜 발화율을 높인다. 즉 소자 자체가 \"광세기→주파수\" 인코더 역할을 하는 뉴런이다. "
     "MNIST 픽셀의 흑백 명암을 광세기로 보고, 소자의 실측 변환 곡선을 통해 주파수로 바꾼 뒤 "
     "스파이크열로 인코딩한다.")
code_block("MNIST 명암 ─[실측 소자 곡선]→ 주파수 ─[Poisson]→ 스파이크 ─[SNN+surrogate]→ 분류")

# 2
h("2. 이론적 배경", 1)
h("2.1 LIF (Leaky Integrate-and-Fire) 뉴런", 2)
para("이산 시간 LIF 막전위 업데이트:  U[t+1] = β·U[t] + I[t] − S[t]·θ")
bullet("β = exp(−dt/τ) : 막전위 누설 계수")
bullet("I[t] : 입력 전류 (광응답 소자에서는 광전류 I_ph)")
bullet("θ : 발화 임계전압, 도달 시 스파이크 S=1 발생 후 리셋")
para("정전류 I 주입 시 발화 조건은 I/(1−β) > θ 이므로 rheobase 전류 = (1−β)·θ. 이보다 큰 "
     "전류일수록 발화 주파수가 증가하는 단조 f–I 곡선이 형성된다 → \"광세기→주파수\" 변환의 물리적 근거.")
h("2.2 Rate Coding (발화율 부호화)", 2)
para("아날로그 값을 스파이크열로 바꾸는 방식. 본 프로젝트는 Poisson rate coding을 사용한다. "
     "각 타임스텝마다 발화 확률 p로 베르누이 시행을 수행하며, p는 측정 주파수 f에 비례한다.")
h("2.3 Surrogate Gradient", 2)
para("스파이크(계단) 함수는 미분이 0/∞이라 역전파가 불가능하다. 순전파는 계단 함수를 쓰되 "
     "역전파 시에만 매끄러운 함수(fast sigmoid)의 기울기로 근사하여 BPTT로 학습한다. "
     "snnTorch의 surrogate.fast_sigmoid() 사용.")

# 3
h("3. 시스템 구조", 1)
table(["단계", "파일", "역할"],
      [["Step 1", "01_photo_neuron.py", "(참고) 수식 기반 광-뉴런 모델, 변환 곡선 검증"],
       ["Step 2", "02_measured_encoder.py", "실측 곡선 기반 grayscale→frequency→Poisson 인코더"],
       ["Step 3", "03_train.py", "MNIST + 인코더 → 2층 SNN 학습/평가"]])

# 4
h("4. 코드 상세", 1)
h("4.1 Step 1 — 광-뉴런 소자 모델 (01_photo_neuron.py)", 2)
para("소자를 하나의 LIF 뉴런으로 모델링한 참고용 구현. 광응답을 포화형(Michaelis-Menten) + "
     "암전류 수식으로 기술한다:  I_ph(Φ) = I_dark + R·Φ^γ/(Φ^γ + Φ_half^γ).")
table(["파라미터", "의미"],
      [["R (responsivity)", "응답도(포화 시 최대 광전류)"],
       ["Φ_half", "반포화 광세기"],
       ["γ (gamma)", "응답 비선형성 (γ<1 sub-linear, =1 포화, >1 super-linear)"],
       ["I_dark", "암전류"]])
para("PhotoLIF 클래스가 이 광전류를 snn.Leaky에 입력해 스파이크를 생성하고, measure_transfer()가 "
     "광세기별 발화 주파수를 측정한다. (출력: fig1 변환 곡선, fig2 막전위·래스터). "
     "실제 학습에는 Step 2의 실측 기반 인코더를 사용한다.")

h("4.2 Step 2 — 측정 데이터 기반 인코더 (02_measured_encoder.py)", 2)
para("실제 소자 측정 보정 곡선을 불러와 MNIST 명암을 실측 주파수로 변환하고 Poisson 스파이크열로 "
     "인코딩한다. 이 프로젝트의 핵심 모듈.")
code_block("g∈[0,1] ─(운용범위)→ I_phys=g·intensity_white ─(실측 보간)→ f(g)[Hz]\n"
           "         ─(정규화)→ p(g)=p_target·f(g)/f_ref ─(spikegen.rate)→ Poisson 스파이크")
para("설계상 두 가지가 중요하다.", bold=True)
bullet("운용 범위(intensity_white): 흰색(g=1)이 대응할 광세기. 측정 곡선이 50 kHz로 강하게 포화하므로 "
       "전 구간(0~9.43)을 쓰면 g>0.2 픽셀이 포화에 몰려 대비가 뭉개진다. 응답 구간(기본 2.0 mW/cm²)만 사용해 대비 확보.")
bullet("auto-scale 정규화: p=p_target·f/f_ref 형태라 dt·절대주파수가 상쇄되고 상대적 명암 정보만 확률로 남는다.")
para("아래는 MNIST 숫자 '5'의 인코딩 예시 (원본 명암 → 측정 주파수 맵 → 스파이크 수).")
image("fig4_mnist_encoding.png", 6.4, "그림. MNIST 명암 → 실측 주파수 → 스파이크 인코딩")

h("4.3 Step 3 — SNN 학습 (03_train.py)", 2)
code_block("spk_in (T,B,784) → FC(784→256) → LIF → FC(256→10) → LIF → spk_out (T,B,10)")
bullet("예측 = 출력 뉴런별 총 스파이크 수가 최대인 클래스 (rate readout)")
bullet("손실: SF.ce_count_loss (출력 스파이크 수에 대한 cross-entropy)")
bullet("최적화: Adam(lr=1e-3), surrogate gradient = fast_sigmoid, β=0.9, num_steps=50, batch=128")
bullet("Step 2의 MeasuredEncoder를 importlib로 재사용")
para("데이터는 MNIST. CPU 시간을 고려해 기본은 서브셋(train 6000 / test 1000); 전체는 "
     "CFG[\"n_train\"]=None으로 전환. 출력: snn_mnist.pt(가중치), fig6_training.png(곡선).")

# 5
h("5. 측정 데이터 및 인코딩 설계", 1)
para("측정 파일: C:\\Claude\\AI Cap\\data\\characteristic_curve.csv "
     "(컬럼 intensity_mW_cm2, frequency_Hz / 0~9.43 mW/cm², 0~50 kHz, 포화형).")
para("intensity_white = 2.0 mW/cm² 설정에서 명암별 인코딩(발화 확률 분포):")
table(["g (명암)", "I [mW/cm²]", "f [Hz]", "p (per-step)"],
      [["0.00", "0.0", "0", "0.00"],
       ["0.25", "0.5", "~2,250", "~0.05"],
       ["0.50", "1.0", "~18,100", "~0.40"],
       ["0.75", "1.5", "~31,800", "~0.70"],
       ["1.00", "2.0", "~40,800", "0.90"]])
image("fig3_measured_transfer.png", 6.4, "그림. 좌: 소자 실측 특성 / 우: grayscale→주파수·확률 인코딩")

# 6
h("6. 실험 환경", 1)
bullet("OS: Windows 11 Pro / Python 3.10 (전용 venv)")
bullet("주요 패키지: torch 2.5.1+cpu, snntorch 0.9.4, torchvision 0.20.1+cpu, pandas, matplotlib")
bullet("환경 주의: 전역 Miniconda torch 2.12.0이 c10.dll 초기화 실패(WinError 1114)로 import 불가 → "
       "전용 venv(.venv)에 검증된 CPU torch 설치로 해결")
bullet("matplotlib 한글: Malgun Gothic 폰트 적용")

# 7
h("7. 학습 결과 (100 epoch)", 1)
para("서브셋(train 6000 / test 1000), batch 128, lr 1e-3, num_steps 50, hidden 256 조건에서 "
     "100 에포크 학습한 결과:")
table(["epoch", "loss", "train acc", "test acc"],
      [["1", "1.312", "76.8%", "70.4%"],
       ["5", "0.122", "97.7%", "92.4%"],
       ["10", "0.034", "98.9%", "93.1%"],
       ["20", "0.004", "99.9%", "94.6%"],
       ["30", "0.001", "100.0%", "94.8%"],
       ["50", "0.000", "100.0%", "94.7%"],
       ["60", "0.004", "99.9%", "95.1% (최고)"],
       ["100", "0.000", "100.0%", "94.6%"]])
image("fig6_training.png", 6.4, "그림. 100 에포크 학습 손실(좌)·정확도(우) 곡선")
para("관찰", bold=True)
bullet("약 10 에포크 내 빠르게 수렴(test 93%), 이후 완만히 상승하여 최고 test 95.1%(epoch 60).")
bullet("train 정확도는 ~30 에포크에 100%, 손실 0으로 포화 → 과적합. 이후 test는 94~95%에서 정체. "
       "서브셋(6000장)만 사용한 결과로 예상된 거동.")
bullet("epoch ~53 부근에서 손실이 일시적으로 튀고 정확도가 잠깐 하락 후 즉시 회복하는 transient "
       "(Adam 최적화 중 일시적 불안정). 학습 안정성에는 영향 없음.")
para("결론: 소자 실측 곡선 기반 인코딩 + SNN 학습 파이프라인이 정상 동작하며, 제한된 서브셋 "
     "조건에서 test 95% 수준 분류 성능을 달성. 성능 향상에는 전체 데이터 학습, 정규화"
     "(dropout/weight decay), early stopping(epoch 60 부근)이 필요하다.", bold=False)

# 8
h("8. 고찰 및 향후 과제", 1)
bullet("인코딩 충실도: 소자 실측 변환 곡선을 인코더에 직접 반영해 실제 하드웨어 동작에 근접한 부호화 구현.")
bullet("운용 범위(intensity_white) 선택이 명암 대비(=분류 성능)에 핵심적.")
bullet("향후: 전체 MNIST 학습 및 GPU 가속 / 측정 곡선의 시간 응답(시정수)까지 반영한 LIF 동역학(전류 역산 모드) / "
       "다른 데이터셋·소자별 변환 곡선 비교.")

out = os.path.join(HERE, "SNN_보고서.docx")
doc.save(out)
print("[saved]", out)
